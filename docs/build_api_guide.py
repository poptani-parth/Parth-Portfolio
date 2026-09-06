from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parent / "Portfolio_API_Guide.docx"
NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F6F8"
GRAY = "5A6573"
WHITE = "FFFFFF"
BLACK = "111827"
CONTENT_WIDTH = 6.5

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)

def set_cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths):
    table.autofit = False
    total = int(sum(widths) * 1440)
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(int(width * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(width * 1440)))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margin(cell)

def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def add_text(doc, text, style=None, before=0, after=6, size=None, color=None, bold=None, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p

def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    shade(cell, "172033")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(text.splitlines()):
        if index: p.add_run("\n")
        r = p.add_run(line)
        set_run_font(r, size=8.2, color=WHITE, name="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_run_font(r, size=9.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10, color=BLACK)

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10, color=BLACK)

def endpoint_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_widths(table, [1.0, 2.95, 2.55])
    headers = ["Method", "Endpoint", "Purpose"]
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value); set_run_font(r, size=9, color=WHITE, bold=True)
    for method, endpoint, purpose in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, [method, endpoint, purpose]):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value); set_run_font(r, size=8.6, color=BLACK, bold=(cell == cells[0]))
            shade(cell, LIGHT_GRAY if len(table.rows) % 2 == 0 else WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_heading(doc, title, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=15 if level == 1 else 11.5, color=NAVY if level == 1 else BLUE, bold=True)
    return p

def api_block(doc, method, path, description, body=None, notes=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(method + " "); set_run_font(r, size=10, color=BLUE, bold=True, name="Consolas")
    r = p.add_run(path); set_run_font(r, size=10, color=BLACK, bold=True, name="Consolas")
    add_text(doc, description, after=3, size=9.6, color=BLACK)
    if body: add_code(doc, body)
    if notes: add_note(doc, "HOW TO USE", notes)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10)
for name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[name].font.name = "Aptos Display"
    styles[name]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    styles[name]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = header.add_run("PORTFOLIO BACKEND  |  API REFERENCE")
set_run_font(hr, size=8, color=GRAY, bold=True)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Portfolio API Guide  •  Local base URL: http://localhost:8080")
set_run_font(fr, size=8, color=GRAY)

# Cover
add_text(doc, "PORTFOLIO BACKEND", size=10, color=BLUE, bold=True, before=32, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "API Reference &\nIntegration Guide", size=28, color=NAVY, bold=True, before=0, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "A practical guide to every public and admin endpoint, JWT authentication, and Postman workflows.", size=12, color=GRAY, before=0, after=34, align=WD_ALIGN_PARAGRAPH.CENTER)
meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(meta, [1.55, 4.95])
for i, (label, value) in enumerate([
    ("Audience", "Portfolio frontend developer and administrator"),
    ("API base URL", "http://localhost:8080"),
    ("Authentication", "Bearer JWT for /api/admin/** (except login, refresh, and password reset)")
]):
    shade(meta.cell(i, 0), NAVY); shade(meta.cell(i, 1), LIGHT_BLUE)
    for j, value2 in enumerate([label, value]):
        p = meta.cell(i, j).paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value2); set_run_font(r, size=9.5, color=WHITE if j == 0 else BLACK, bold=(j == 0))
doc.add_page_break()

add_heading(doc, "1. Getting started")
add_text(doc, "This guide describes the Spring Boot portfolio API currently exposed at the local base URL. JSON is used for request and response bodies unless an endpoint is marked multipart/form-data.", size=10)
add_heading(doc, "Base URL and required headers", 2)
add_code(doc, "Base URL: http://localhost:8080\nContent-Type: application/json\nAdmin calls: Authorization: Bearer {{access_token}}")
add_note(doc, "POSTMAN", "Create an environment variable named base_url = http://localhost:8080. After login, save accessToken from the response as access_token and refreshToken as refresh_token.")
add_heading(doc, "Quick admin workflow", 2)
for item in [
    "POST /api/admin/auth/login with your username and password.",
    "Copy accessToken into the Authorization > Bearer Token tab for admin requests.",
    "When the access token expires, POST /api/admin/auth/refresh with refresh_token. This returns a new access and refresh token; replace both stored values.",
    "Use the protected admin endpoints to create, update, delete, upload, or review content."
]: numbered(doc, item)
add_note(doc, "SECURITY", "Access tokens last 15 minutes by default. Refresh tokens are rotated and become unusable after a successful refresh. Do not place either token in a public repository or browser local storage.")

add_heading(doc, "2. Authentication and account APIs")
endpoint_table(doc, [
    ("POST", "/api/admin/auth/login", "Authenticate the admin and receive access + refresh tokens"),
    ("POST", "/api/admin/auth/refresh", "Rotate a refresh token and obtain a new token pair"),
    ("POST", "/api/admin/auth/password-reset", "Email a one-time password-reset link"),
    ("POST", "/api/admin/auth/password-reset/confirm", "Set a new password using the emailed token"),
    ("POST", "/api/admin/auth/change-password", "Change the password while authenticated")
])
api_block(doc, "POST", "/api/admin/auth/login", "Public endpoint. Username is case-sensitive. A successful response includes tokenType = Bearer and expiresIn in seconds.",
'''{
  "username": "admin",
  "password": "your-strong-password"
}''', "In Postman, choose Body > raw > JSON. A 401 means the account is missing, inactive, or the stored BCrypt hash does not match the supplied password.")
add_code(doc, '''{
  "accessToken": "eyJ...",
  "refreshToken": "eyJ...",
  "tokenType": "Bearer",
  "expiresIn": 900
}''')
api_block(doc, "POST", "/api/admin/auth/refresh", "Public endpoint. Submit the most recently issued refresh token. Use the returned pair immediately and discard the old pair.",
'''{
  "refreshToken": "{{refresh_token}}"
}''')
api_block(doc, "POST", "/api/admin/auth/password-reset", "Public endpoint. Always returns 204 to avoid revealing whether the username exists. The reset email is sent to ADMIN_RESET_EMAIL.",
'''{
  "username": "admin"
}''')
api_block(doc, "POST", "/api/admin/auth/password-reset/confirm", "Public endpoint. The token is single-use and expires after 15 minutes. The new password must be 12-128 characters with upper-case, lower-case, number, and symbol.",
'''{
  "username": "admin",
  "token": "token-from-reset-email",
  "newPassword": "NewStrongPassword!2026"
}''')
api_block(doc, "POST", "/api/admin/auth/change-password", "Admin JWT required. Revokes every existing refresh token after a successful change.",
'''{
  "currentPassword": "current-password",
  "newPassword": "NewStrongPassword!2026"
}''')

add_heading(doc, "3. Public portfolio APIs")
add_text(doc, "Public endpoints do not need a JWT. They expose active/public portfolio content only.", size=10)
endpoint_table(doc, [
    ("GET", "/api/profile", "Public profile"),
    ("GET", "/api/skills?category={category}", "Active skills, grouped by category; category is optional"),
    ("GET", "/api/skills/{id}", "One skill by MongoDB id"),
    ("GET", "/api/experience", "Active work history, newest first"),
    ("GET", "/api/education", "Active education entries"),
    ("GET", "/api/projects?category={category}&tech={tag}", "Active projects; both filters are optional"),
    ("GET", "/api/projects/{id}", "One project by MongoDB id"),
    ("GET", "/api/media?usage={usage}", "Active media; optional usage such as PROFILE, PROJECT, CERTIFICATE"),
    ("POST", "/api/contact", "Submit contact form; rate limited")
])
api_block(doc, "GET", "/api/projects?category=Web&tech=Spring", "Filters are optional and can be used independently. URL-encode spaces and special characters in query values.")
api_block(doc, "POST", "/api/contact", "Public contact form. The server validates input, sanitizes text content, stores the message, and sends configured mail.",
'''{
  "name": "Jane Doe",
  "email": "jane@example.com",
  "subject": "Portfolio enquiry",
  "message": "Hello, I would like to discuss a project."
}''', "Validation: name <= 100, email <= 150, subject <= 200, message <= 3000 characters. The default rate limit is 5 submissions per IP per 10 minutes.")

add_heading(doc, "4. Admin content APIs")
add_text(doc, "Every endpoint in this section requires the ADMIN role. Set Authorization: Bearer {{access_token}} in Postman. Use a 24-character MongoDB id for {id} path variables.", size=10)
add_heading(doc, "Profile", 2)
endpoint_table(doc, [("GET", "/api/admin/profile", "Read admin profile"), ("POST", "/api/admin/profile", "Create profile"), ("PUT", "/api/admin/profile", "Update profile")])
api_block(doc, "POST / PUT", "/api/admin/profile", "Create or update the profile. Use the same JSON shape for either operation.",
'''{
  "name": "Your Name",
  "heroTitle": "Backend Developer",
  "heroSubtitle": "Building reliable applications",
  "bio": "Short biography text",
  "email": "you@example.com",
  "phone": "+94 77 000 0000",
  "location": "Colombo, Sri Lanka",
  "resumeUrl": "https://example.com/resume.pdf",
  "githubUrl": "https://github.com/your-account",
  "linkedinUrl": "https://linkedin.com/in/your-account",
  "profileImageUrl": "https://example.com/profile.jpg",
  "active": true
}''')
add_heading(doc, "Skills", 2)
endpoint_table(doc, [("GET", "/api/admin/skills", "List all skills"), ("GET", "/api/admin/skills/{id}", "Read one skill"), ("POST", "/api/admin/skills", "Create skill"), ("PUT", "/api/admin/skills/{id}", "Update skill"), ("DELETE", "/api/admin/skills/{id}", "Delete skill"), ("PATCH", "/api/admin/skills/reorder", "Change ordering")])
api_block(doc, "POST / PUT", "/api/admin/skills[/ {id}]", "Use POST without an id to create; use PUT with an id to update.", '''{
  "name": "Spring Boot",
  "category": "Backend",
  "displayOrder": 1,
  "active": true
}''')
api_block(doc, "PATCH", "/api/admin/skills/reorder", "Send every changed id/order pair in an array.", '''[
  { "id": "64f000000000000000000001", "order": 0 },
  { "id": "64f000000000000000000002", "order": 1 }
]''')
add_heading(doc, "Projects", 2)
endpoint_table(doc, [("GET", "/api/admin/projects", "List all projects"), ("GET", "/api/admin/projects/{id}", "Read one project"), ("POST", "/api/admin/projects", "Create project"), ("PUT", "/api/admin/projects/{id}", "Update project"), ("DELETE", "/api/admin/projects/{id}", "Delete project")])
api_block(doc, "POST / PUT", "/api/admin/projects[/ {id}]", "Use arrays for technologies and features. The text fields are stored as sanitized plain text.", '''{
  "title": "Portfolio Backend",
  "description": "Spring Boot portfolio API.",
  "category": "Web",
  "technologies": ["Java", "Spring Boot", "MongoDB"],
  "githubUrl": "https://github.com/example/repo",
  "liveUrl": "https://example.com",
  "imageUrl": "https://example.com/project.jpg",
  "role": "Backend Developer",
  "features": ["JWT authentication", "Admin content management"],
  "displayOrder": 1,
  "featured": true,
  "active": true
}''')

add_heading(doc, "5. Experience and education APIs")
add_heading(doc, "Experience", 2)
endpoint_table(doc, [("GET", "/api/admin/experience", "List all experience entries"), ("GET", "/api/admin/experience/{id}", "Read one experience entry"), ("POST", "/api/admin/experience", "Create entry"), ("PUT", "/api/admin/experience/{id}", "Update entry"), ("DELETE", "/api/admin/experience/{id}", "Delete entry")])
api_block(doc, "POST / PUT", "/api/admin/experience[/ {id}]", "Dates use ISO format YYYY-MM-DD. When current is true, endDate must be null. When current is false, endDate is required and cannot precede startDate.", '''{
  "company": "Example Company",
  "position": "Software Engineer",
  "location": "Colombo",
  "employmentType": "Full-time",
  "description": "Built and maintained backend services.",
  "startDate": "2024-01-01",
  "endDate": null,
  "current": true,
  "displayOrder": 0,
  "active": true
}''')
add_heading(doc, "Education", 2)
endpoint_table(doc, [("GET", "/api/admin/education", "List all education entries"), ("GET", "/api/admin/education/{id}", "Read one education entry"), ("POST", "/api/admin/education", "Create entry"), ("PUT", "/api/admin/education/{id}", "Update entry"), ("DELETE", "/api/admin/education/{id}", "Delete entry")])
api_block(doc, "POST / PUT", "/api/admin/education[/ {id}]", "endDate is optional. displayOrder must be zero or greater.", '''{
  "institution": "Example University",
  "degree": "BSc Computer Science",
  "field": "Software Engineering",
  "startDate": "2020-01-01",
  "endDate": "2024-01-01",
  "grade": "First Class",
  "description": "Relevant coursework and achievements.",
  "displayOrder": 0,
  "active": true
}''')

add_heading(doc, "6. Media and contact-message APIs")
add_heading(doc, "Media", 2)
endpoint_table(doc, [("POST", "/api/admin/media/upload", "Upload image or PDF"), ("GET", "/api/admin/media", "List all media"), ("DELETE", "/api/admin/media/{id}", "Delete media and stored file")])
api_block(doc, "POST", "/api/admin/media/upload", "Use Body > form-data in Postman, not raw JSON. Set the file row type to File. Allowed image types: JPEG, PNG, WEBP (max 5 MB); PDF (max 10 MB).", None,
"Required form-data keys: file (File), usage (Text: profile, project, resume, certificate, or general). Optional: altText (max 300) and caption (max 500).")
add_heading(doc, "Contact messages", 2)
endpoint_table(doc, [("GET", "/api/admin/contact", "List submitted messages; optional ?status="), ("GET", "/api/admin/contact/{id}", "Read one message"), ("PATCH", "/api/admin/contact/{id}/read", "Mark read"), ("PATCH", "/api/admin/contact/{id}/archive", "Archive")])
add_note(doc, "STATUS FILTER", "Use the ContactMessageStatus enum value accepted by the backend, for example ?status=NEW. No request body is needed for the read/archive PATCH calls.")

add_heading(doc, "7. Responses, errors, and troubleshooting")
endpoint_table(doc, [("200", "OK", "Successful read, update, login, or refresh"), ("201", "Created", "Successful create or upload"), ("204", "No Content", "Successful delete, password action, or reset request"), ("400", "Bad Request", "Validation, malformed JSON, or invalid input"), ("401", "Unauthorized", "Missing/invalid JWT or invalid login credentials"), ("403", "Forbidden", "JWT is valid but does not have ADMIN"), ("404", "Not Found", "Requested resource does not exist"), ("429", "Too Many Requests", "Login or contact rate limit exceeded"), ("500", "Server Error", "Unexpected backend error")])
add_heading(doc, "Common Postman problems", 2)
bullet(doc, "401 on login: Confirm the MongoDB password field is a BCrypt hash for the exact password entered and that active is true. Use the password-reset flow to set a fresh BCrypt password if needed.")
bullet(doc, "401 on an admin call: Log in again and paste the new accessToken into Authorization > Bearer Token. Do not use the refresh token for normal API calls.")
bullet(doc, "403 on an admin call: The token does not contain the ADMIN role. Verify the admin user role is ADMIN, then log in again.")
bullet(doc, "400 from JSON endpoints: Choose raw > JSON, set Content-Type: application/json, remove extra braces, and follow required field limits.")
bullet(doc, "429: Wait for the Retry-After duration before retrying. Repeated login attempts extend the lockout window.")
bullet(doc, "CORS error in a browser: Set FRONTEND_ORIGIN to the exact frontend scheme, domain, and port; Postman is not affected by CORS.")
add_heading(doc, "Security notes", 2)
add_text(doc, "All admin write actions are audited. The API rejects MongoDB operator-shaped JSON fields and strips HTML from stored text. Never commit .env, JWT secrets, mail credentials, database URIs, tokens, or passwords.", size=10)

doc.core_properties.title = "Portfolio Backend API Reference"
doc.core_properties.subject = "Endpoint and Postman integration guide"
doc.core_properties.author = "Portfolio Backend"
doc.save(OUT)
print(OUT)
